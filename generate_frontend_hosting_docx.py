import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0052CC"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(33, 37, 41)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_callout(doc, title, text, box_type="info"):
    colors = {
        "info": {"bg": "EBF5FB", "border": "2980B9", "text_color": "1B4F72"},
        "warning": {"bg": "FEF9E7", "border": "F39C12", "text_color": "7D6608"},
        "success": {"bg": "EAFAF1", "border": "27AE60", "text_color": "145A32"}
    }
    cfg = colors.get(box_type, colors["info"])
    
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, cfg["bg"])
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{cfg['border']}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    r_title = p.add_run(f"📌 {title}: ")
    r_title.bold = True
    r_title.font.name = "Khmer OS Siemreap"
    r_title.font.size = Pt(10)
    
    r_text = p.add_run(text)
    r_text.font.name = "Khmer OS Siemreap"
    r_text.font.size = Pt(10)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_frontend_hosting_doc(output_path):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Khmer OS Siemreap'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(40, 40, 40)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("មគ្គុទ្ទេសក៍ការដាក់ដំណើរការ FRONTEND លើ AWS UBUNTU SERVER\n(React Vite SPA + Nginx Web Server on AWS EC2: 34.232.147.247)")
    run_title.font.name = 'Khmer OS Muol Light'
    run_title.font.size = Pt(14)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle / Info
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(10)
    run_sub = sub_p.add_run("AWS Server IP: 34.232.147.247 | Port: 80 (HTTP) | React 18 + Vite + TailwindCSS")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(9.5)
    run_sub.bold = True
    run_sub.font.color.rgb = RGBColor(0, 102, 153)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(10)
    p_div_run = p_div.add_run("―" * 55)
    p_div_run.font.color.rgb = RGBColor(200, 200, 200)
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECTION 1: OVERVIEW & SPECS
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("១. ព័ត៌មានលម្អិតអំពីប្រព័ន្ធ Frontend (Frontend System Specs)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph("Frontend ត្រូវបានបង្កើតឡើងដោយប្រើ React Vite Single Page Application (SPA) និងបំពាក់ដោយ Nginx Web Server ដែលមានល្បឿនលឿន និងស៊ីធនធាន RAM តិចបំផុត៖")
    p.paragraph_format.space_after = Pt(6)

    # Table Specs
    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.3)
    tbl.columns[1].width = Inches(4.2)
    
    spec_data = [
        ("AWS Server Public IP", "34.232.147.247"),
        ("Web Frontend Port", "Port 80 (HTTP Standard Web Port)"),
        ("Technology Stack", "React 18, Vite, TailwindCSS, Axios, Leaflet Map, Face API"),
        ("Web Server Engine", "Nginx Alpine (Gzip Compression, SPA Fallback, Cache-Control)"),
        ("API Routing", "Reverse Proxy /api/ -> Backend Spring Boot API")
    ]

    for i, (k, v) in enumerate(spec_data):
        c0 = tbl.cell(i, 0)
        c1 = tbl.cell(i, 1)
        set_cell_background(c0, "EBF1F5" if i % 2 == 0 else "F4F6F8")
        set_cell_background(c1, "FFFFFF" if i % 2 == 0 else "F9F9F9")
        set_cell_margins(c0, top=50, bottom=50, left=90, right=90)
        set_cell_margins(c1, top=50, bottom=50, left=90, right=90)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.name = "Segoe UI"
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(v)
        r1.font.name = "Khmer OS Siemreap"
        r1.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # AWS Inbound Rules Callout
    add_callout(doc, "ការកំណត់ AWS Security Group Inbound Rules", "នៅលើ AWS Console -> EC2 -> Security Groups សូមបើក Ports ដូចខាងក្រោម៖\n• Port 80 (HTTP): បើក 0.0.0.0/0 ដើម្បីឱ្យអ្នកប្រើប្រាស់អាចបើក Web ពីគ្រប់ទីកន្លែង\n• Port 22 (SSH): សម្រាប់ចូលបញ្ជា Terminal ពីចម្ងាយ\n• Port 8080 (Custom TCP): សម្រាប់ Backend Spring Boot API", "warning")

    # SECTION 2: SSH LOGIN
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)
    r = h2.add_run("២. របៀបភ្ជាប់ SSH ទៅកាន់ AWS Ubuntu Server")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("បើក Terminal (ឬ PowerShell) នៅលើកុំព្យូទ័ររបស់អ្នក រួចដំណើរការបញ្ជា៖")
    add_code_block(doc, 'chmod 400 "C:/path/to/your-key.pem"\nssh -i "C:/path/to/your-key.pem" ubuntu@34.232.147.247')

    # SECTION 3: 1-COMMAND DEPLOYMENT
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)
    r = h3.add_run("៣. ការដាក់ដំណើរការ Frontend ដោយស្វ័យប្រវត្តិ (1-Command Fast Deploy)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("នៅពេលចូលដល់ Ubuntu Server សូមចូលទៅកាន់ Folder Frontend រួច Run Script តែមួយគត់៖")
    add_code_block(doc, '# ១. ចូលទៅកាន់ Folder Frontend\ncd Hr_chomnan/frontend\n\n# ២. ផ្ដល់សិទ្ធិដំណើរការ deploy.sh\nchmod +x deploy.sh\n\n# ៣. ដំណើរការ Deploy ភ្លាមៗ\n./deploy.sh')

    add_callout(doc, "តើ deploy.sh ធ្វើអ្វីខ្លះដោយស្វ័យប្រវត្តិ?", "១. ត្រួតពិនិត្យ និងដំឡើង Docker + Docker Compose បើមិនទាន់មានលើ Server\n២. Build Production Web Bundle ដោយស្វ័យប្រវត្តិតាមរយៈ Multi-stage Dockerfile\n៣. បង្កើត Container Nginx និងដំណើរការលើ Port 80 ដោយស្វ័យប្រវត្តិ\n៤. Health Check ពិនិត្យមើលថាតើ Web Page អាចបើកបានជោគជ័យឬនៅ", "success")

    # SECTION 4: MANUAL DEPLOYMENT STEPS
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(6)
    r = h4.add_run("៤. វិធីដាក់ដំណើរការដោយដៃ (Manual Step-by-Step Deploy)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("ជម្រើសទី ១៖ ប្រើ Docker Compose (វិធីដែលងាយស្រួល និងសុវត្ថិភាពបំផុត)៖")
    add_code_block(doc, '# ចូលទៅកាន់ Frontend Directory\ncd Hr_chomnan/frontend\n\n# Build និង Run Container លើ Port 80\ndocker compose up -d --build\n\n# ពិនិត្យ Status\ndocker compose ps')

    doc.add_paragraph("ជម្រើសទី ២៖ Build Static Files រួចដាក់លើ Nginx ផ្ទាល់លើ Host (Standalone)៖")
    add_code_block(doc, '# ១. ដំឡើង Node.js 20 និង Nginx\nsudo apt-get update\nsudo apt-get install -y nginx nodejs npm\n\n# ២. Install Dependencies និង Build React App\ncd Hr_chomnan/frontend\nnpm install\nnpm run build\n\n# ៣. Copy Files ទៅកាន់ /var/www/html\nsudo cp -r dist/* /var/www/html/\nsudo systemctl restart nginx')

    # SECTION 5: MANAGEMENT COMMANDS
    h5 = doc.add_heading(level=1)
    h5.paragraph_format.space_before = Pt(10)
    h5.paragraph_format.space_after = Pt(6)
    r = h5.add_run("៥. តារាងពាក្យបញ្ជាគ្រប់គ្រង Frontend សំខាន់ៗ (Useful Commands)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    # Table Commands
    tbl_cmd = doc.add_table(rows=6, cols=2)
    tbl_cmd.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cmd.autofit = False
    tbl_cmd.columns[0].width = Inches(2.4)
    tbl_cmd.columns[1].width = Inches(4.1)
    
    cmd_data = [
        ("មើល Logs របស់ Nginx / Frontend", "docker logs -f hr_attendance_frontend"),
        ("Restart Frontend Server", "docker compose restart"),
        ("បិទ Frontend Server (Stop)", "docker compose down"),
        ("Update កូដថ្មីរួច Re-build", "git pull && docker compose up -d --build"),
        ("ពិនិត្យ Container Status", "docker compose ps"),
        ("ពិនិត្យមើល RAM/CPU Usage", "docker stats hr_attendance_frontend")
    ]

    for i, (desc, cmd) in enumerate(cmd_data):
        c0 = tbl_cmd.cell(i, 0)
        c1 = tbl_cmd.cell(i, 1)
        set_cell_background(c0, "EBF1F5" if i % 2 == 0 else "F4F6F8")
        set_cell_background(c1, "FFFFFF" if i % 2 == 0 else "F9F9F9")
        set_cell_margins(c0, top=45, bottom=45, left=80, right=80)
        set_cell_margins(c1, top=45, bottom=45, left=80, right=80)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(desc)
        r0.font.name = "Khmer OS Siemreap"
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(cmd)
        r1.font.name = "Consolas"
        r1.font.size = Pt(9)
        r1.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # SECTION 6: URL ACCESS
    h6 = doc.add_heading(level=1)
    h6.paragraph_format.space_before = Pt(10)
    h6.paragraph_format.space_after = Pt(6)
    r = h6.add_run("៦. អាសយដ្ឋានចូលប្រើប្រាស់ Web Application")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("បន្ទាប់ពី Deploy ចប់ អ្នកអាចបើក Web Browser (Google Chrome / Safari) រួចចូលតាម Link ខាងក្រោម៖")

    # Table Web URLs
    tbl_url = doc.add_table(rows=4, cols=2)
    tbl_url.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_url.autofit = False
    tbl_url.columns[0].width = Inches(2.2)
    tbl_url.columns[1].width = Inches(4.3)
    
    url_data = [
        ("Web Frontend (Home)", "http://34.232.147.247"),
        ("Admin & Staff Login", "http://34.232.147.247/login"),
        ("Kiosk Face Attendance", "http://34.232.147.247/kiosk"),
        ("Employee Management", "http://34.232.147.247/employees")
    ]

    for i, (name, link) in enumerate(url_data):
        c0 = tbl_url.cell(i, 0)
        c1 = tbl_url.cell(i, 1)
        set_cell_background(c0, "EBF1F5" if i % 2 == 0 else "F4F6F8")
        set_cell_background(c1, "FFFFFF" if i % 2 == 0 else "F9F9F9")
        set_cell_margins(c0, top=40, bottom=40, left=70, right=70)
        set_cell_margins(c1, top=40, bottom=40, left=70, right=70)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(name)
        r0.bold = True
        r0.font.name = "Segoe UI"
        r0.font.size = Pt(9)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(link)
        r1.font.name = "Consolas"
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor(0, 102, 204)

    # Footer note
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(16)
    r_f = p_foot.add_run("― ចប់មគ្គុទ្ទេសក៍ការដាក់ដំណើរការ FRONTEND លើ AWS UBUNTU ―\nServer IP: 34.232.147.247 | Port 80 (HTTP) | HR Chomnan System")
    r_f.font.name = "Khmer OS Siemreap"
    r_f.font.size = Pt(9)
    r_f.font.color.rgb = RGBColor(120, 120, 120)
    doc.save(output_path)
    print(f"Document successfully generated at: {output_path}")

if __name__ == "__main__":
    build_frontend_hosting_doc(r"d:\project\Hr_chomnan\AWS_Ubuntu_Frontend_Hosting_Guide.docx")
