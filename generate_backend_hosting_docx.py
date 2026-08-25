import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0052CC"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(33, 37, 41)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_callout(doc, title, text, box_type="info"):
    colors = {
        "info": {"bg": "EBF5FB", "border": "2980B9", "text_color": "1B4F72"},
        "warning": {"bg": "FEF9E7", "border": "F39C12", "text_color": "7D6608"},
        "success": {"bg": "EAFAF1", "border": "27AE60", "text_color": "145A32"}
    }
    cfg = colors.get(box_type, colors["info"])
    
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, cfg["bg"])
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{cfg['border']}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    r_title = p.add_run(f"📌 {title}: ")
    r_title.bold = True
    r_title.font.name = "Khmer OS Siemreap"
    r_title.font.size = Pt(10)
    
    r_text = p.add_run(text)
    r_text.font.name = "Khmer OS Siemreap"
    r_text.font.size = Pt(10)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_backend_hosting_doc(output_path):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Khmer OS Siemreap'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(40, 40, 40)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("មគ្គុទ្ទេសក៍ការដាក់ដំណើរការ BACKEND លើ AWS UBUNTU SERVER\n(Spring Boot 3.4 + PostgreSQL 16 on AWS EC2)")
    run_title.font.name = 'Khmer OS Muol Light'
    run_title.font.size = Pt(14)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle / Info
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(10)
    run_sub = sub_p.add_run("Spring Boot (Port 8080) | PostgreSQL (Port 5432) | Docker Installation Included")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(9.5)
    run_sub.bold = True
    run_sub.font.color.rgb = RGBColor(0, 102, 153)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(10)
    p_div_run = p_div.add_run("―" * 55)
    p_div_run.font.color.rgb = RGBColor(200, 200, 200)
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECTION 1: OVERVIEW & SPECS
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("១. ព័ត៌មានលម្អិតអំពីប្រព័ន្ធ Backend (System Overview)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph("ប្រព័ន្ធ Backend ត្រូវបានរៀបចំឡើងយ៉ាងស្អាត ដោយដំណើរការលើ Container ដាច់ដោយឡែកពីគ្នា មានសុវត្ថិភាព និងងាយស្រួលក្នុងការ Host លើ AWS EC2 Ubuntu Server៖")
    p.paragraph_format.space_after = Pt(6)

    # Table Specs
    tbl = doc.add_table(rows=6, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.3)
    tbl.columns[1].width = Inches(4.2)
    
    spec_data = [
        ("Architecture", "Docker Compose Multi-Container Setup"),
        ("Spring Boot Backend", "Java 21 (Eclipse Temurin) - Port 8080 (Context: /api)"),
        ("Database Engine", "PostgreSQL 16 Alpine - Port 5432 (Internal Docker Network)"),
        ("Database Name", "employee_attendance_db (Persistent Volume: postgres_data)"),
        ("Security & Auth", "JWT Stateless Authentication (24-hour expiration)"),
        ("Timezone", "Asia/Phnom_Penh (Cambodia Standard Time GMT+7)")
    ]

    for i, (k, v) in enumerate(spec_data):
        c0 = tbl.cell(i, 0)
        c1 = tbl.cell(i, 1)
        set_cell_background(c0, "EBF1F5" if i % 2 == 0 else "F4F6F8")
        set_cell_background(c1, "FFFFFF" if i % 2 == 0 else "F9F9F9")
        set_cell_margins(c0, top=50, bottom=50, left=90, right=90)
        set_cell_margins(c1, top=50, bottom=50, left=90, right=90)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.name = "Segoe UI"
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(v)
        r1.font.name = "Khmer OS Siemreap"
        r1.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # AWS Inbound Rules Callout
    add_callout(doc, "ការកំណត់ AWS Security Group Inbound Rules", "នៅលើ AWS Console -> EC2 -> Security Groups សូមបើក Ports ដូចខាងក្រោមជាចាំបាច់៖\n• Port 22 (SSH): សម្រាប់ចូលបញ្ជា Terminal (My IP ឬ 0.0.0.0/0)\n• Port 8080 (Custom TCP): សម្រាប់ Backend API / Mobile App (0.0.0.0/0)\n• Port 80 (HTTP): ប្រសិនបើចង់បើក Web Frontend ជាមួយគ្នា (0.0.0.0/0)", "warning")

    # SECTION 2: SSH LOGIN
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)
    r = h2.add_run("២. របៀបភ្ជាប់ SSH ទៅកាន់ AWS Ubuntu Server")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("បើក Terminal (ឬ PowerShell) នៅលើកុំព្យូទ័ររបស់អ្នក រួចដំណើរការបញ្ជា៖")
    add_code_block(doc, 'chmod 400 "C:/path/to/your-key.pem"\nssh -i "C:/path/to/your-key.pem" ubuntu@YOUR_SERVER_IP')

    # SECTION 3: DOCKER INSTALLATION
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)
    r = h3.add_run("៣. របៀបដំឡើង Docker និង Docker Compose លើ Ubuntu Server")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("ប្រសិនបើ Ubuntu Server របស់អ្នកមិនទាន់មាន Docker សូម Copy & Paste បញ្ជាខាងក្រោមដើម្បីដំឡើង Docker Engine និង Docker Compose ផ្លូវការ (Official Docker Engine)៖")

    doc.add_paragraph("ក. ដំឡើង Docker ពេញលេញដោយស្វ័យប្រវត្តិ (Official Script):")
    add_code_block(doc, '# ១. Update Package List & Install Prereqs\nsudo apt-get update -y\nsudo apt-get install -y ca-certificates curl gnupg lsb-release\n\n# ២. បន្ថែម Official Docker GPG Key & Repository\nsudo mkdir -p /etc/apt/keyrings\ncurl -fsSL https://download.docker.com/linux/ubuntu/gpg | sudo gpg --dearmor -o /etc/apt/keyrings/docker.gpg --yes\necho "deb [arch=$(dpkg --print-architecture) signed-by=/etc/apt/keyrings/docker.gpg] https://download.docker.com/linux/ubuntu $(lsb_release -cs) stable" | sudo tee /etc/apt/sources.list.d/docker.list > /dev/null\n\n# ៣. ដំឡើង Docker Engine + CLI + Containerd + Docker Compose\nsudo apt-get update -y\nsudo apt-get install -y docker-ce docker-ce-cli containerd.io docker-buildx-plugin docker-compose-plugin docker-compose\n\n# ៤. ផ្ដល់សិទ្ធិឱ្យ User បច្ចុប្បន្នប្រើ Docker ដោយមិនបាច់វាយ sudo\nsudo usermod -aG docker $USER\nsudo systemctl enable --now docker\nnewgrp docker')

    doc.add_paragraph("ខ. ផ្ទៀងផ្ទាត់ការដំឡើង Docker (Verification):")
    add_code_block(doc, 'docker --version\ndocker compose version')

    add_callout(doc, "លទ្ធផលជោគជ័យនៃការដំឡើង", "Docker Version នឹងបង្ហាញ: Docker version 27.x ឬថ្មីជាងនេះ\nDocker Compose នឹងបង្ហាញ: Docker Compose version v2.x ឬថ្មីជាងនេះ", "success")

    # SECTION 4: 1-COMMAND DEPLOYMENT
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(6)
    r = h4.add_run("៤. ការដាក់ដំណើរការ Backend ដោយស្វ័យប្រវត្តិ (1-Command Fast Deploy)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("នៅពេលចូលដល់ Ubuntu Server សូមចូលទៅកាន់ Folder Backend រួច Run Script តែមួយគត់៖")
    add_code_block(doc, '# ១. ចូលទៅកាន់ Folder Backend\ncd Hr_chomnan/backend\n\n# ២. ផ្ដល់សិទ្ធិដំណើរការ Scripts\nchmod +x deploy.sh backup_db.sh\n\n# ៣. ដំណើរការ Deploy ភ្លាមៗ\n./deploy.sh')

    add_callout(doc, "តើ deploy.sh ធ្វើអ្វីខ្លះដោយស្វ័យប្រវត្តិ?", "១. ត្រួតពិនិត្យ និងដំឡើង Docker + Docker Compose បើមិនទាន់មាន\n២. បង្កើត File .env ដោយស្វ័យប្រវត្តិ និងបង្កើត Secure Random JWT Key\n៣. Build Docker Image សម្រាប់ Spring Boot (Java 21) ដោយប្រើ Multi-stage Caching\n៤. បង្កើត Container PostgreSQL 16 និង Spring Boot Backend ឱ្យដំណើរការ\n៥. ត្រួតពិនិត្យសុខភាព (Health Check) នៃ API រហូតដល់ដំណើរការ ១០០%", "success")

    # SECTION 5: MANUAL DEPLOYMENT STEPS
    h5 = doc.add_heading(level=1)
    h5.paragraph_format.space_before = Pt(10)
    h5.paragraph_format.space_after = Pt(6)
    r = h5.add_run("៥. វិធីដាក់ដំណើរការដោយដៃ (Manual Step-by-Step Deploy)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("ក. បង្កើត និងកំណត់ File .env សម្រាប់ Backend៖")
    add_code_block(doc, 'cd Hr_chomnan/backend\ncp .env.example .env\nnano .env')

    doc.add_paragraph("ខ. ចាប់ផ្ដើមដំណើរការ Backend & PostgreSQL ជាមួយ Docker Compose៖")
    add_code_block(doc, '# Build និង Run Containers ក្នុង Background\ndocker compose up -d --build\n\n# ពិនិត្យមើល Status របស់ Containers\ndocker compose ps')

    # SECTION 6: API TESTING & VERIFICATION
    h6 = doc.add_heading(level=1)
    h6.paragraph_format.space_before = Pt(10)
    h6.paragraph_format.space_after = Pt(6)
    r = h6.add_run("៦. ការតេស្តសាកល្បង API (Testing & Health Check Endpoints)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("បន្ទាប់ពី Deploy រួចរាល់ អ្នកអាចតេស្ត API តាមវិធីដូចខាងក្រោម៖")

    # Table API Endpoints
    tbl_api = doc.add_table(rows=4, cols=3)
    tbl_api.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_api.autofit = False
    tbl_api.columns[0].width = Inches(1.8)
    tbl_api.columns[1].width = Inches(1.2)
    tbl_api.columns[2].width = Inches(3.5)
    
    api_data = [
        ("Health Check", "GET", "http://YOUR_SERVER_IP:8080/api/health"),
        ("Kiosk Settings", "GET", "http://YOUR_SERVER_IP:8080/api/kiosk-settings"),
        ("Auth Login", "POST", "http://YOUR_SERVER_IP:8080/api/auth/login"),
        ("Face Attendance", "POST", "http://YOUR_SERVER_IP:8080/api/attendance/face-kiosk")
    ]

    for i, (name, method, url) in enumerate(api_data):
        c0 = tbl_api.cell(i, 0)
        c1 = tbl_api.cell(i, 1)
        c2 = tbl_api.cell(i, 2)
        set_cell_background(c0, "EBF1F5" if i % 2 == 0 else "F4F6F8")
        set_cell_background(c1, "E8F8F5" if method == "GET" else "FEF9E7")
        set_cell_background(c2, "FFFFFF" if i % 2 == 0 else "F9F9F9")
        set_cell_margins(c0, top=40, bottom=40, left=70, right=70)
        set_cell_margins(c1, top=40, bottom=40, left=70, right=70)
        set_cell_margins(c2, top=40, bottom=40, left=70, right=70)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(name)
        r0.bold = True
        r0.font.name = "Segoe UI"
        r0.font.size = Pt(9)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(method)
        r1.bold = True
        r1.font.name = "Consolas"
        r1.font.size = Pt(9)
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(url)
        r2.font.name = "Consolas"
        r2.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.add_paragraph("តេស្តតាម Terminal លើ Server (Curl Command)៖")
    add_code_block(doc, '# តេស្តសុខភាព API\ncurl -i http://localhost:8080/api/kiosk-settings\n\n# តេស្តពីម៉ាស៊ីនខាងក្រៅ (កុំព្យូទ័ររបស់អ្នក)\ncurl -i http://YOUR_SERVER_IP:8080/api/kiosk-settings')

    # SECTION 7: MANAGEMENT COMMANDS
    h7 = doc.add_heading(level=1)
    h7.paragraph_format.space_before = Pt(10)
    h7.paragraph_format.space_after = Pt(6)
    r = h7.add_run("៧. តារាងពាក្យបញ្ជាគ្រប់គ្រង Server សំខាន់ៗ (Useful Commands)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    # Table Commands
    tbl_cmd = doc.add_table(rows=7, cols=2)
    tbl_cmd.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cmd.autofit = False
    tbl_cmd.columns[0].width = Inches(2.4)
    tbl_cmd.columns[1].width = Inches(4.1)
    
    cmd_data = [
        ("មើល Logs របស់ Backend ផ្ទាល់", "docker logs -f hr_attendance_backend"),
        ("មើល Logs របស់ PostgreSQL", "docker logs -f hr_attendance_postgres"),
        ("Restart Server ឡើងវិញ", "docker compose restart"),
        ("បិទ Server (Stop)", "docker compose down"),
        ("Update កូដថ្មីរួច Re-build", "git pull && docker compose up -d --build"),
        ("ចូលទៅក្នុង PostgreSQL DB CLI", "docker exec -it hr_attendance_postgres psql -U postgres -d employee_attendance_db"),
        ("ពិនិត្យការប្រើប្រាស់ RAM/CPU", "docker stats")
    ]

    for i, (desc, cmd) in enumerate(cmd_data):
        c0 = tbl_cmd.cell(i, 0)
        c1 = tbl_cmd.cell(i, 1)
        set_cell_background(c0, "EBF1F5" if i % 2 == 0 else "F4F6F8")
        set_cell_background(c1, "FFFFFF" if i % 2 == 0 else "F9F9F9")
        set_cell_margins(c0, top=45, bottom=45, left=80, right=80)
        set_cell_margins(c1, top=45, bottom=45, left=80, right=80)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(desc)
        r0.font.name = "Khmer OS Siemreap"
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(cmd)
        r1.font.name = "Consolas"
        r1.font.size = Pt(9)
        r1.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # SECTION 8: BACKUP & RESTORE
    h8 = doc.add_heading(level=1)
    h8.paragraph_format.space_before = Pt(10)
    h8.paragraph_format.space_after = Pt(6)
    r = h8.add_run("៨. ការ Backup និង Restore ទិន្នន័យ Database")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("ក. ការ Backup ទិន្នន័យដោយប្រើ Script៖")
    add_code_block(doc, './backup_db.sh\n# ទិន្នន័យនឹងត្រូវរក្សាទុកក្នុង Folder: ./backups/db_backup_YYYYMMDD_HHMMSS.sql.gz')

    doc.add_paragraph("ខ. ការ Restore ទិន្នន័យចូល Database វិញ៖")
    add_code_block(doc, '# ពន្លា និងបញ្ជូនទិន្នន័យចូល Database វិញ\ngunzip < backups/db_backup_YYYYMMDD_HHMMSS.sql.gz | docker exec -i hr_attendance_postgres psql -U postgres -d employee_attendance_db')

    # Footer note
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(16)
    r_f = p_foot.add_run("― ចប់មគ្គុទ្ទេសក៍ការដាក់ដំណើរការ BACKEND លើ AWS UBUNTU ―\nSpring Boot 3.4.2 + PostgreSQL 16 | HR Chomnan System")
    r_f.font.name = "Khmer OS Siemreap"
    r_f.font.size = Pt(9)
    r_f.font.color.rgb = RGBColor(120, 120, 120)

    try:
        doc.save(output_path)
        print(f"Document successfully generated at: {output_path}")
    except PermissionError:
        alt_path = output_path.replace(".docx", "_v2.docx")
        doc.save(alt_path)
        print(f"File locked by Word. Saved alternative to: {alt_path}")

if __name__ == "__main__":
    build_backend_hosting_doc(r"d:\project\Hr_chomnan\AWS_Ubuntu_Backend_Hosting_Guide.docx")
