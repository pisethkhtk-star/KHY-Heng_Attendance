import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F4F6F8")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0052CC"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(33, 37, 41)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_callout(doc, title, text, box_type="info"):
    colors = {
        "info": {"bg": "EBF5FB", "border": "2980B9", "text_color": "1B4F72"},
        "warning": {"bg": "FEF9E7", "border": "F39C12", "text_color": "7D6608"},
        "success": {"bg": "EAFAF1", "border": "27AE60", "text_color": "145A32"}
    }
    cfg = colors.get(box_type, colors["info"])
    
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, cfg["bg"])
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{cfg['border']}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    r_title = p.add_run(f"📌 {title}: ")
    r_title.bold = True
    r_title.font.name = "Khmer OS Siemreap"
    r_title.font.size = Pt(10)
    
    r_text = p.add_run(text)
    r_text.font.name = "Khmer OS Siemreap"
    r_text.font.size = Pt(10)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_deployment_word_doc(output_path):
    doc = docx.Document()
    
    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Khmer OS Siemreap'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(40, 40, 40)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("មគ្គុទ្ទេសក៍ការដាក់ដំណើរការ DOCKER ពេញលេញ\n(Fullstack Docker Deployment Guide: Backend + Frontend + PostgreSQL)")
    run_title.font.name = 'Khmer OS Muol Light'
    run_title.font.size = Pt(15)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0, 51, 102)

    # Subtitle / Info
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(12)
    run_sub = sub_p.add_run("AWS EC2 Ubuntu Server: 100.56.149.110 | Spring Boot 3.4 (Java 21) + React Vite (Nginx) + PostgreSQL 16")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(9.5)
    run_sub.bold = True
    run_sub.font.color.rgb = RGBColor(0, 102, 153)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(10)
    p_div_run = p_div.add_run("―" * 55)
    p_div_run.font.color.rgb = RGBColor(200, 200, 200)
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # SECTION 1: ARCHITECTURE OVERVIEW
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("១. ស្ថាបត្យកម្មប្រព័ន្ធ Fullstack Docker (Architecture Overview)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph("ប្រព័ន្ធទាំងមូលត្រូវបានចងភ្ជាប់គ្នាដោយស្វ័យប្រវត្តិតាមរយៈ Docker Compose ដែលមាន ៣ សេវាកម្មសំខាន់ៗដំណើរការលើ Container ដាច់ដោយឡែកពីគ្នា៖")
    p.paragraph_format.space_after = Pt(6)

    # Table Info
    tbl = doc.add_table(rows=6, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(2.2)
    tbl.columns[1].width = Inches(4.3)
    
    info_data = [
        ("AWS Server Public IP", "100.56.149.110"),
        ("1. Frontend Container", "React Vite + Nginx (Port 80 -> http://100.56.149.110) & Proxy /api/ ទៅ Backend"),
        ("2. Backend Container", "Spring Boot 3.4.2 (Port 8080 -> 8080) ដំណើរការ API ទាំងអស់"),
        ("3. Database Container", "PostgreSQL 16 Alpine (Port 5432) ភ្ជាប់តាមរយៈ Network ផ្ទៃក្នុង"),
        ("Internal Docker Network", "hr-network (Bridge Network សម្រាប់ Container ទំនាក់ទំនងគ្នា)"),
        ("Database Volume", "postgres_data (រក្សាទុកទិន្នន័យលើ Server មិនឱ្យបាត់បង់)")
    ]

    for i, (k, v) in enumerate(info_data):
        c0 = tbl.cell(i, 0)
        c1 = tbl.cell(i, 1)
        set_cell_background(c0, "EBF1F5" if i % 2 == 0 else "F4F6F8")
        set_cell_background(c1, "FFFFFF" if i % 2 == 0 else "F9F9F9")
        set_cell_margins(c0, top=50, bottom=50, left=90, right=90)
        set_cell_margins(c1, top=50, bottom=50, left=90, right=90)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.name = "Segoe UI"
        r0.font.size = Pt(9.5)
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(v)
        r1.font.name = "Khmer OS Siemreap"
        r1.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # AWS Inbound Rules
    add_callout(doc, "ការកំណត់ AWS Security Group Inbound Rules", "នៅលើ AWS Console -> EC2 -> Security Groups សូមបើក Ports ដូចខាងក្រោម៖\n• Port 22 (SSH): សម្រាប់ចូលបញ្ជា Terminal ពីចម្ងាយ\n• Port 80 (HTTP): សម្រាប់ចូលប្រើ Web Frontend (http://100.56.149.110)\n• Port 8080 (Custom TCP): សម្រាប់ Backend API ផ្ទាល់ ឬ Mobile App\n• Port 443 (HTTPS): ប្រសិនបើរៀបចំ SSL ទៅថ្ងៃក្រោយ", "info")

    # SECTION 2: DOCKER SETUP ON UBUNTU
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("២. ជំហានដំឡើង Docker & Docker Compose លើ Ubuntu Server")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    r = p.add_run("ក. ភ្ជាប់ចូលទៅកាន់ AWS EC2 តាមរយៈ SSH:")
    r.bold = True
    add_code_block(doc, "ssh -i /path/to/your-key.pem ubuntu@100.56.149.110")

    p = doc.add_paragraph()
    r = p.add_run("ខ. Update Ubuntu និងដំឡើង Official Docker Engine:")
    r.bold = True
    add_code_block(doc, """# ១. Update System Packages
sudo apt update && sudo apt upgrade -y
sudo apt install -y ca-certificates curl gnupg lsb-release git

# ២. បន្ថែម Docker Official GPG Key និង Repository
sudo mkdir -p /etc/apt/keyrings
curl -fsSL https://download.docker.com/linux/ubuntu/gpg | sudo gpg --dearmor -o /etc/apt/keyrings/docker.gpg

echo "deb [arch=$(dpkg --print-architecture) signed-by=/etc/apt/keyrings/docker.gpg] https://download.docker.com/linux/ubuntu $(lsb_release -cs) stable" | sudo tee /etc/apt/sources.list.d/docker.list > /dev/null

# ៣. ដំឡើង Docker និង Docker Compose Plugin
sudo apt update
sudo apt install -y docker-ce docker-ce-cli containerd.io docker-buildx-plugin docker-compose-plugin

# ៤. អនុញ្ញាតឱ្យ User ubuntu ដំណើរការ Docker ដោយមិនបាច់វាយ sudo
sudo usermod -aG docker $USER
newgrp docker

# ៥. ពិនិត្យមើល Version បញ្ជាក់
docker --version
docker compose version""")

    # SECTION 3: PROJECT SETUP & .ENV
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("៣. ទាញយកគម្រោង និងរៀបចំ File បរិស្ថាន (.env)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    r = p.add_run("ក. Clone គម្រោងទាំងមូលចូលទៅកាន់ Server:")
    r.bold = True
    add_code_block(doc, """cd ~
git clone <your-repository-url> Hr_chomnan
cd Hr_chomnan""")

    p = doc.add_paragraph()
    r = p.add_run("ខ. បង្កើត File .env នៅក្នុង Root Directory នៃគម្រោង:")
    r.bold = True
    add_code_block(doc, """nano .env""")
    
    p = doc.add_paragraph("ចម្លងទិន្នន័យកំណត់ខាងក្រោមចូលក្នុង `.env`:")
    add_code_block(doc, """# Database Configuration
POSTGRES_DB=employee_attendance_db
POSTGRES_USER=postgres
POSTGRES_PASSWORD=SecurePassword2026!@#
POSTGRES_PORT=5432

# Backend Configuration
BACKEND_PORT=8080
JWT_SECRET=404E635266556A586E3272357538782F413F4428472B4B6250645367566B5970
TZ=Asia/Phnom_Penh

# Frontend Configuration
FRONTEND_PORT=80
VITE_API_BASE_URL=""
""")

    add_callout(doc, "ចំណាំលើ VITE_API_BASE_URL", "ដោយសារ Frontend Nginx Container ត្រូវបានរៀបចំ Reverse Proxy ផ្ទេរ /api/ ទៅកាន់ Backend រួចជាស្រេច យើងកំណត់ VITE_API_BASE_URL=\"\" (ទទេ) ដើម្បីឱ្យ Frontend ហៅ /api/ ដោយស្វ័យប្រវត្តិតាម Same-Origin Domain ដោយមិនបារម្ភបញ្ហា CORS ឡើយ។", "success")

    # SECTION 4: DOCKER COMPOSE FILE STRUCTURE
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("៤. រចនាសម្ព័ន្ធ docker-compose.yml សម្រាប់ Backend + Frontend + DB")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph("ឯកសារ `docker-compose.yml` នៅក្នុង Folder ដើម (`~/Hr_chomnan/docker-compose.yml`) មានទម្រង់ដូចខាងក្រោម៖")
    add_code_block(doc, """services:
  postgres:
    image: postgres:16-alpine
    container_name: hr_attendance_postgres
    restart: unless-stopped
    environment:
      POSTGRES_DB: ${POSTGRES_DB:-employee_attendance_db}
      POSTGRES_USER: ${POSTGRES_USER:-postgres}
      POSTGRES_PASSWORD: ${POSTGRES_PASSWORD:-postgres}
      TZ: Asia/Phnom_Penh
    ports:
      - "${POSTGRES_PORT:-5432}:5432"
    volumes:
      - postgres_data:/var/lib/postgresql/data
    healthcheck:
      test: ["CMD-SHELL", "pg_isready -U ${POSTGRES_USER:-postgres} -d ${POSTGRES_DB:-employee_attendance_db}"]
      interval: 10s
      timeout: 5s
      retries: 5
    networks:
      - hr-network

  backend:
    build:
      context: ./backend
      dockerfile: Dockerfile
    image: hr-chomnan-backend:latest
    container_name: hr_attendance_backend
    restart: unless-stopped
    environment:
      DATABASE_URL: jdbc:postgresql://postgres:5432/${POSTGRES_DB:-employee_attendance_db}
      DB_USERNAME: ${POSTGRES_USER:-postgres}
      DB_PASSWORD: ${POSTGRES_PASSWORD:-postgres}
      JWT_SECRET: ${JWT_SECRET:-your-super-secret-jwt-key-change-this}
      TZ: Asia/Phnom_Penh
    ports:
      - "${BACKEND_PORT:-8080}:8080"
    depends_on:
      postgres:
        condition: service_healthy
    networks:
      - hr-network

  frontend:
    build:
      context: ./frontend
      dockerfile: Dockerfile
      args:
        VITE_API_BASE_URL: ${VITE_API_BASE_URL:-}
    image: hr-chomnan-frontend:latest
    container_name: hr_attendance_frontend
    restart: unless-stopped
    environment:
      TZ: Asia/Phnom_Penh
    ports:
      - "${FRONTEND_PORT:-80}:80"
    depends_on:
      - backend
    networks:
      - hr-network

networks:
  hr-network:
    driver: bridge

volumes:
  postgres_data:
    driver: local""")

    # SECTION 5: ONE COMMAND DEPLOYMENT
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("៥. ដំណើរការ Deploy ទាំងមូលក្នុងពេលតែមួយ (1-Command Deploy)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph("គ្រាន់តែដំណើរការ Command តែ ១ ខាងក្រោម Docker នឹងធ្វើការទាញយក DB, Build Java Backend និង Build React Frontend ព្រមទាំងចងភ្ជាប់ Network ដោយស្វ័យប្រវត្តិ៖")
    add_code_block(doc, """# ១. ចូលទៅកាន់ Folder គម្រោង
cd ~/Hr_chomnan

# ២. បញ្ជាឱ្យ Build និង Start សេវាកម្មទាំងអស់ក្នុង Background
docker compose up -d --build

# ៣. ពិនិត្យមើលស្ថានភាពដំណើរការនៃ Containers
docker compose ps""")

    p = doc.add_paragraph()
    r = p.add_run("🔹 ពិនិត្យមើល Logs បញ្ជាក់ការចាប់ផ្តើម:")
    r.bold = True
    add_code_block(doc, """# មើល Logs របស់ Backend
docker compose logs -f backend

# មើល Logs របស់ Frontend
docker compose logs -f frontend

# មើល Logs របស់ Database
docker compose logs -f postgres""")

    add_callout(doc, "សាកល្បងចូលប្រើប្រាស់", "បន្ទាប់ពីដំណើរការចប់ អ្នកអាចបើក Browser ចូលតាមអាសយដ្ឋាន៖\n• Web Frontend: http://100.56.149.110 (ដំណើរការលើ Port 80)\n• Backend Direct API: http://100.56.149.110:8080/api/auth/login", "success")

    # SECTION 6: DOCKER MANAGEMENT COMMANDS
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("៦. តារាងពាក្យបញ្ជាគ្រប់គ្រង Docker (Docker Commands Cheatsheet)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    # Table Commands
    tbl_cmd = doc.add_table(rows=8, cols=2)
    tbl_cmd.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cmd.autofit = False
    tbl_cmd.columns[0].width = Inches(3.4)
    tbl_cmd.columns[1].width = Inches(3.1)
    
    cmd_data = [
        ("docker compose up -d --build", "Build និង Start Containers ទាំងអស់ (Backend + Frontend + DB)"),
        ("docker compose up -d --build backend", "Re-build និង Update តែសេវាកម្ម Backend មួយគត់"),
        ("docker compose up -d --build frontend", "Re-build និង Update តែសេវាកម្ម Frontend មួយគត់"),
        ("docker compose ps", "ពិនិត្យមើលស្ថានភាព Containers (Status, Ports)"),
        ("docker compose logs -f --tail=50 backend", "តាមដាន Logs របស់ Backend ចំនួន 50 ជួរចុងក្រោយ"),
        ("docker compose restart backend frontend", "Restart សេវាកម្ម Backend និង Frontend ឡើងវិញ"),
        ("docker compose down", "បិទ Containers ទាំងអស់ (ទិន្នន័យ DB នៅគង់វង្សក្នុង Volume)"),
        ("docker system prune -af", "សម្អាត Docker Images ចាស់ៗដែលមិនប្រើ ដើម្បីសន្សំទំហំ Disk")
    ]

    for i, (k, v) in enumerate(cmd_data):
        c0 = tbl_cmd.cell(i, 0)
        c1 = tbl_cmd.cell(i, 1)
        set_cell_background(c0, "EBF1F5" if i % 2 == 0 else "F4F6F8")
        set_cell_background(c1, "FFFFFF" if i % 2 == 0 else "F9F9F9")
        set_cell_margins(c0, top=50, bottom=50, left=80, right=80)
        set_cell_margins(c1, top=50, bottom=50, left=80, right=80)
        
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        r0 = p0.add_run(k)
        r0.font.name = "Consolas"
        r0.font.size = Pt(9)
        r0.bold = True
        
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        r1 = p1.add_run(v)
        r1.font.name = "Khmer OS Siemreap"
        r1.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # SECTION 7: UPDATE CODE WORKFLOW
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("៧. របៀប Update កូដថ្មីឡើងលើ Server (CI/CD Update Workflow)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph("នៅពេលអ្នកកែប្រែកូដថ្មីរួចហើយ Push ឡើង GitHub/GitLab អ្នកគ្រាន់តែធ្វើដូចខាងក្រោម៖")
    add_code_block(doc, """cd ~/Hr_chomnan

# ១. ទាញយកកូដថ្មីចុងក្រោយ
git pull origin main

# ២. ករណីកែប្រែកូដទាំង Backend និង Frontend
docker compose up -d --build

# (ឬ) ករណីកែតែ Backend មួយប៉ុណ្ណោះ:
# docker compose up -d --build backend

# (ឬ) ករណីកែតែ Frontend មួយប៉ុណ្ណោះ:
# docker compose up -d --build frontend""")

    # SECTION 8: MOBILE APP CONFIG
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("៨. ការកំណត់ Mobile App (Flutter) ឱ្យភ្ជាប់មកកាន់ Server IP")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph("នៅក្នុង Project Mobile App `mobile_app/lib/core/constants/api_config.dart` សូមកំណត់ដូចខាងក្រោម៖")
    add_code_block(doc, """class ApiConfig {
  // កំណត់ Server IP របស់ AWS
  static const String serverHost = '100.56.149.110';
  
  // កំណត់ Port 8080 (ឬទុកទទេ '' ប្រសិនបើហៅតាម Port 80)
  static const String serverPort = '8080';
  
  // ប្រើ HTTP
  static const bool useHttps = false;
}""")

    # SECTION 9: DATABASE BACKUP
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("៩. ការ Backup និង Restore Database តាមរយៈ Docker")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    r = p.add_run("ក. ទាញយក Backup ទិន្នន័យចេញជាឯកសារ .sql:")
    r.bold = True
    add_code_block(doc, "docker exec -t hr_attendance_postgres pg_dump -U postgres employee_attendance_db > ~/backup_$(date +%F_%H%M%S).sql")

    p = doc.add_paragraph()
    r = p.add_run("ខ. Restore ទិន្នន័យចូល Database វិញ:")
    r.bold = True
    add_code_block(doc, "cat ~/backup_file.sql | docker exec -i hr_attendance_postgres psql -U postgres -d employee_attendance_db")

    # SECTION 10: SWAP MEMORY
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r = h1.add_run("១០. ការបង្កើត Swap Memory (ការពារការខ្វះ RAM ពេល Build លើ AWS EC2)")
    r.font.name = "Khmer OS Muol Light"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph("ដោយសារការ Build Java (Gradle) និង React (Node.js) ត្រូវការ RAM ច្រើន ប្រសិនបើ Server EC2 មាន RAM ត្រឹម 1GB/2GB សូមបង្កើត Swap 2GB ជាមុនសិន៖")
    add_code_block(doc, """sudo fallocate -l 2G /swapfile
sudo chmod 600 /swapfile
sudo mkswap /swapfile
sudo swapon /swapfile
echo '/swapfile none swap sw 0 0' | sudo tee -a /etc/fstab
free -h""")

    # Footer
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(20)
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_foot.add_run("― ចប់មគ្គុទ្ទេសក៍ការដាក់ដំណើរការ FULLSTACK DOCKER ―\nServer IP: 100.56.149.110 | HR Chomnan Fullstack System")
    r_f.font.size = Pt(9.5)
    r_f.font.color.rgb = RGBColor(120, 120, 120)

    doc.save(output_path)
    print(f"Successfully created: {output_path}")

if __name__ == "__main__":
    out_file = r"d:\project\Hr_chomnan\AWS_Ubuntu_Fullstack_Docker_Deployment_Guide.docx"
    try:
        build_deployment_word_doc(out_file)
    except PermissionError:
        out_file = r"d:\project\Hr_chomnan\AWS_Ubuntu_Fullstack_Docker_Deployment_Guide_New.docx"
        build_deployment_word_doc(out_file)
